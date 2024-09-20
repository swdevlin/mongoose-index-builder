import argparse
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm

from webindex import generate_web_index

INTRO_PAGES = 4

LEFT_ALIGNMENT = 0

GROUP_CORRECTIONS = {
	'Augment': 'Augmentations',
	'Augments': 'Augmentations',
	'Augmentation': 'Augmentations',
	'Robot': 'Robots',
	'Drone': 'Drones',
	'Ship': 'Ships',
	'Vehicle': 'Vehicles',
	'Armour': 'Personal Protection',
}

# Bulk correct entries with the incorrect type
TYPE_CORRECTIONS = {
	'Armour': 'Personal Protection',
	'Career': 'Careers',
	'Drone': 'Drones',
	# Not sure person needs to be its own category; adding them to setting for the time being
	'Person': 'Setting',
	'Skill': 'Skills',
	'Robot': 'Robots',
	'Ship': 'Ships',
	'Sophant': 'Sophants',
	'small craft': 'Small Craft',
	'Small craft': 'Small Craft',
	'Vehicle': 'Vehicles',
}


def has_page_break(para):
	for run in para.runs:
		if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
			return True
	return False


def delete_existing_entries(document):
	page_count = 1
	position_found = False
	for para in document.paragraphs:
		if position_found:
			p = para._element
			p.getparent().remove(p)
		else:
			if has_page_break(para):
				page_count += 1
				position_found = page_count >= INTRO_PAGES


def add_type_paragraph(document, type_name):
	document.add_page_break()
	paragraph = document.add_paragraph()
	run = paragraph.add_run(type_name)
	run.bold = True
	run.font.size = Pt(12)
	paragraph.alignment = LEFT_ALIGNMENT
	return paragraph


def add_page_text(paragraph, topic):
	paragraph.add_run('\t')

	books = sorted(topic['entries'].keys())
	for i, book in enumerate(books):
		pages = topic['entries'][book]
		p = ','.join(str(page['page']) for page in pages)
		entry_text = paragraph.add_run(f"{book} {p}")
		entry_text.font.size = Pt(8)
		if i < len(books) - 1:
			comma = paragraph.add_run(", ")
			comma.font.size = Pt(8)
	paragraph.alignment = LEFT_ALIGNMENT


def add_index_line(document, topic):
	paragraph = document.add_paragraph()

	subject = topic['topic']
	subject_text = paragraph.add_run(subject)
	if len(subject) > 35:
		subject_text.font.size = Pt(9)
	else:
		subject_text.font.size = Pt(10)

	if len(topic['entries']) > 0:
		paragraph.paragraph_format.tab_stops.add_tab_stop(
			Cm(8.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
		)
		add_page_text(paragraph, topic)

	for child_key in sorted(topic['children'].keys()):
		child = topic['children'][child_key]
		paragraph = document.add_paragraph()

		subject = child['topic']
		subject_text = paragraph.add_run(subject)
		subject_text.font.size = Pt(10)

		paragraph.paragraph_format.tab_stops.add_tab_stop(
			Cm(8.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
		)
		paragraph.paragraph_format.left_indent = Cm(0.5)

		add_page_text(paragraph, child)


def create_traveller_index(topics, filename):
	document = Document(filename)
	delete_existing_entries(document)
	section = document.add_section(0)
	sectPr = section._sectPr
	cols = sectPr.xpath("./w:cols")[0]
	cols.set(qn('w:num'), '2')

	last_type = None
	for index, topic in enumerate(topics):
		if topic['type'] != last_type:
			add_type_paragraph(document, topic['type'])
			last_type = topic['type']

		add_index_line(document, topic)

	document.save(filename)


def add_topic(key, row, topics):
	if not key in topics:
		topics[key] = {
			'children': {},
			'entries': {},
			'topic': row['Topic'],
			'type': row['Type'],
		}
	book = row['Document']
	entry = topics[key]['entries']
	if not book in entry:
		entry[book] = []
	entry[book].append({
		"page": row['Page'],
		"primary": row['Primary'] != 'No',
	})


def parse_topics(source, topics):
	df = pd.read_csv(source, delimiter='\t')
	df = df.replace({pd.NA: None, pd.NaT: None, float('nan'): None})
	for index, row in df.iterrows():
		for key in TYPE_CORRECTIONS:
			if row['Type'] == key:
				row['Type'] = TYPE_CORRECTIONS[key]

		for key in GROUP_CORRECTIONS:
			if row.get('Group', None) == key:
				row['Group'] = GROUP_CORRECTIONS[key]

		subject = row['Topic']
		group = row.get('Group', None)
		if group:
			group_key = (row['Type'], group)
			if group_key not in topics:
				topics[group_key] = {
					'children': {},
					'topic': group,
					'type': row['Type'],
					'entries': {},
				}
		else:
			group_key = None

		key = (row['Type'], subject)
		if group_key:
			add_topic(key, row, topics[group_key]['children'])
			if row['Type'] == 'Setting':
				add_topic(key, row, topics)
		else:
			add_topic(key, row, topics)


def main():
	parser = argparse.ArgumentParser(description="Mongoose Traveller Index Generator")
	parser.add_argument('-f', '--file', nargs='+', help='File to process', required=False, default=[])
	parser.add_argument('-d', '--dir', nargs='+', help='Directory to process', required=False, default=[])
	parser.add_argument('-o', '--output', help='Output Word document', required=True)

	args = parser.parse_args()

	topics = {}
	for file in args.file:
		parse_topics(file, topics)

	for directory in args.dir:
		for root, _, filenames in os.walk(directory):
			for filename in filenames:
				if filename.endswith('.tsv'):
					full_path = os.path.join(root, filename)
					parse_topics(full_path, topics)

	sorted_topics = []
	for key in sorted(topics.keys()):
		sorted_topics.append(topics[key])
	create_traveller_index(sorted_topics, args.output)
	generate_web_index(sorted_topics, 'static/webindex.html')


if __name__ == "__main__":
	main()
