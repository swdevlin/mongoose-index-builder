from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm


INTRO_PAGES = 4

LEFT_ALIGNMENT = 0


def has_page_break(para):
	for run in para.runs:
		if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
			return True
	return False


def delete_existing_entries(document):
	page_count = 1
	position_found = False
	for para in document.paragraphs:
		if position_found:
			p = para._element
			p.getparent().remove(p)
		else:
			if has_page_break(para):
				page_count += 1
				position_found = page_count >= INTRO_PAGES


def add_type_paragraph(document, type_name):
	document.add_page_break()
	paragraph = document.add_paragraph()
	run = paragraph.add_run(type_name)
	run.bold = True
	run.font.size = Pt(12)
	paragraph.alignment = LEFT_ALIGNMENT
	return paragraph


def add_page_text(paragraph, topic):
	paragraph.add_run('\t')

	books = sorted(topic['entries'].keys())
	for i, book in enumerate(books):
		pages = topic['entries'][book]
		p = ','.join(str(page['page']) for page in pages)
		entry_text = paragraph.add_run(f"{book} {p}")
		entry_text.font.size = Pt(8)
		if i < len(books) - 1:
			comma = paragraph.add_run(", ")
			comma.font.size = Pt(8)
	paragraph.alignment = LEFT_ALIGNMENT


def split_subject(subject):
	if len(subject) > 35:
		split_pos = subject.rfind(' ', 0, 35)

		part1 = subject[:split_pos].strip()
		part2 = subject[split_pos:].strip()
		return [part1, part2]
	else:
		return [subject]


def add_subject(document, subject, indent):
	parts = split_subject(subject)

	max_length = 35 if indent else 40

	if len(parts[-1]) > max_length:
		font_size = Pt(8)
	elif len(parts[-1]) > max_length-5:
		font_size = Pt(9)
	else:
		font_size = Pt(10)

	if len(parts) == 2:
		paragraph = document.add_paragraph()
		if indent:
			paragraph.paragraph_format.left_indent = Cm(indent)
		subject_text = paragraph.add_run(parts[0])
		subject_text.font.size = font_size
	paragraph = document.add_paragraph()
	subject_text = paragraph.add_run(parts[-1])
	subject_text.font.size = font_size
	if indent:
		paragraph.paragraph_format.left_indent = Cm(indent)
	return paragraph


def add_index_line(document, topic):
	subject = topic['topic']
	paragraph = add_subject(document, subject, None)

	if len(topic['entries']) > 0:
		paragraph.paragraph_format.tab_stops.add_tab_stop(
			Cm(8.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
		)
		add_page_text(paragraph, topic)

	for child_key in sorted(topic['children'].keys()):
		child = topic['children'][child_key]
		subject = child['topic']
		paragraph = add_subject(document, subject, 0.5)

		paragraph.paragraph_format.tab_stops.add_tab_stop(
			Cm(8.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
		)
		add_page_text(paragraph, child)


def create_traveller_index(topics, filename):
	document = Document(filename)
	delete_existing_entries(document)
	section = document.add_section(0)
	sectPr = section._sectPr
	cols = sectPr.xpath("./w:cols")[0]
	cols.set(qn('w:num'), '2')

	last_type = None
	for index, topic in enumerate(topics):
		if topic['type'] != last_type:
			add_type_paragraph(document, topic['type'])
			last_type = topic['type']

		add_index_line(document, topic)

	document.save(filename)
